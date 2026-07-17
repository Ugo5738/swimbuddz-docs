from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

OUT_DIR = Path(__file__).resolve().parent

BRAND_DEEP = colors.HexColor("#063B4F")
BRAND_AQUA = colors.HexColor("#BFEFF2")
BRAND_INK = colors.HexColor("#102A33")
BRAND_MUTED = colors.HexColor("#60727A")
BRAND_BG = colors.HexColor("#F5FBFC")
BRAND_LINE = colors.HexColor("#D8E7EA")


@dataclass(frozen=True)
class Section:
    title: str
    bullets: tuple[str, ...]


@dataclass(frozen=True)
class Variant:
    filename_stem: str
    audience_label: str
    title: str
    subtitle: str
    intro: str
    context: str
    at_glance: tuple[tuple[str, str], ...]
    sections: tuple[Section, ...]
    start_steps: tuple[tuple[str, str], ...]
    next_step: str


GENERIC = Variant(
    filename_stem="SwimBuddz_Intro_to_Water_Pilot_Outline",
    audience_label="DISCUSSION OUTLINE FOR HR REVIEW",
    title="INTRO TO WATER",
    subtitle="A low-friction employee wellness pilot",
    intro=(
        "A company-sponsored first experience that lets employees try swimming safely "
        "before HR or staff commit to a longer programme."
    ),
    context=(
        "Built to answer two practical HR questions: will employees participate, and "
        "will the experience visibly support their wellbeing?"
    ),
    at_glance=(
        ("8-10 employees", "Suggested first group; adjusted to venue and coach coverage."),
        ("2 hours", "Guided, beginner-friendly water-confidence experience."),
        ("Saturday morning", "Low disruption to the working week."),
        ("Lagos partner pool", "Yaba, VI, Ikoyi or Festac; final venue follows participant needs."),
    ),
    sections=(
        Section(
            "Why swimming",
            (
                "Supports physical wellbeing through low-impact cardio, mobility and recovery.",
                "Helps employees decompress beyond the usual office, gym or screen routine.",
                "Builds confidence and water-safety skills they keep beyond work.",
                "Includes adults who avoid water from fear, embarrassment or limited access.",
            ),
        ),
        Section(
            "Why this pilot",
            (
                "Employees experience the value before deciding whether to continue.",
                "Company sponsorship reduces financial friction at the first step.",
                "HR measures real participation instead of relying on announcement interest.",
                "A longer cohort is considered only if genuine demand exists.",
            ),
        ),
        Section(
            "What employees experience",
            (
                "Warm orientation for adults with little or no water confidence.",
                "Guided breathing, floating and safe movement fundamentals.",
                "Small-group coaching with no pressure to perform.",
                "A clear next step only for employees who want to continue.",
            ),
        ),
        Section(
            "What SwimBuddz handles",
            (
                "Interest form and simple pre-session safety screening.",
                "Coach coverage and partner-pool coordination.",
                "Teaching aids where needed: kickboards, pool noodles or float supports.",
                "Employee reminders, session-day logistics and concise HR report.",
            ),
        ),
        Section(
            "What HR receives",
            (
                "Registration, attendance and engagement summary.",
                "Anonymous before/after confidence and wellbeing pulse.",
                "Satisfaction, comments and interest in continuing.",
                "Recommendation to stop, repeat, or progress to a cohort.",
            ),
        ),
        Section(
            "Commercial & privacy",
            (
                "Recommended: the company sponsors the first session so employees pay nothing.",
                "Final pilot fee is confirmed after headcount and venue are agreed.",
                "Personal fears, health notes and skill details stay with the coach.",
                "HR receives only agreed, aggregated feedback.",
            ),
        ),
    ),
    start_steps=(
        ("1 Scope", "20-minute HR call"),
        ("2 Confirm", "Date, venue, headcount"),
        ("3 Launch", "SwimBuddz opens registration"),
        ("4 Report", "Session + report within 3 business days"),
    ),
    next_step=(
        "Suggested next step: a 15-20 minute planning call to confirm headcount, "
        "likely location, sponsorship level and whether the pilot is only for feedback "
        "or a pathway into a 12-week cohort."
    ),
)


NESTUGE = Variant(
    filename_stem="SwimBuddz_Intro_to_Water_Pilot_for_Nestuge",
    audience_label="DISCUSSION OUTLINE FOR NESTUGE",
    title="INTRO TO WATER",
    subtitle="A low-commitment employee wellness pilot",
    intro=(
        "A one-off, company-sponsored weekend session that lets Nestuge test employee "
        "interest before considering a longer swim programme."
    ),
    context=(
        "Designed around four practical realities: employee time, cost, sustained "
        "participation, and Island/Mainland accessibility."
    ),
    at_glance=(
        ("8-10 employees", "Suggested first group; final size follows venue and coach coverage."),
        ("2 hours", "One guided session with no 12-week commitment."),
        ("Saturday morning", "Outside the work week and limited to one session."),
        ("Location-led venue", "Yaba, VI, Ikoyi or Festac, chosen after a participant location check."),
    ),
    sections=(
        Section(
            "Why swimming",
            (
                "Supports physical wellbeing through low-impact cardio, mobility and recovery.",
                "Helps employees decompress beyond the usual office, gym or screen routine.",
                "Builds confidence and water-safety skills they keep beyond work.",
                "Includes adults who avoid water from fear, embarrassment or limited access.",
            ),
        ),
        Section(
            "Why this format",
            (
                "Employees feel the value before weighing a longer programme.",
                "Nestuge measures real demand instead of announcement interest.",
                "A defined start and finish reduces drop-off risk.",
                "A full cohort follows only if enough employees want to continue.",
            ),
        ),
        Section(
            "What employees experience",
            (
                "Warm, beginner-friendly orientation for adults with low water confidence.",
                "Guided breathing, floating and safe movement fundamentals.",
                "Small-group coaching with no pressure to perform.",
                "A clear next step, with no obligation to join the 12-week programme.",
            ),
        ),
        Section(
            "SwimBuddz handles",
            (
                "Interest form, participant location check and pre-session screening.",
                "Pool and coach coordination with a clear session plan.",
                "Teaching aids where needed: kickboards, pool noodles or float supports.",
                "Reminders, session delivery, feedback and HR recommendation.",
            ),
        ),
        Section(
            "Nestuge provides",
            (
                "One People/Operations contact for coordination.",
                "Approval of the date, venue and final quote.",
                "Internal circulation of the registration link and pilot sponsorship.",
                "Decision on whether to stop, repeat, or progress to a cohort.",
            ),
        ),
        Section(
            "Commercial & privacy",
            (
                "Nestuge sponsors the capped first session so employees pay nothing.",
                "Final quote states headcount, venue and pool access.",
                "Personal health, fear and skill details stay with the coach; HR receives aggregated feedback only.",
            ),
        ),
    ),
    start_steps=(
        ("1 Interest check", "Short internal form"),
        ("2 Scope", "Date, areas, headcount, budget"),
        ("3 Confirm & launch", "Quote approved; registration opens"),
        ("4 Run & report", "Session + report within 3 business days"),
    ),
    next_step=(
        "Suggested next step: a short planning call to confirm expected interest, "
        "location mix, budget comfort and whether Nestuge wants feedback only or a "
        "pathway into a 12-week cohort."
    ),
)


def styles():
    base = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle(
            "kicker",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7.6,
            textColor=BRAND_MUTED,
            leading=9,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "title": ParagraphStyle(
            "title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            textColor=BRAND_DEEP,
            leading=30,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11.4,
            textColor=BRAND_INK,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "intro": ParagraphStyle(
            "intro",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.4,
            textColor=BRAND_INK,
            leading=10,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "context": ParagraphStyle(
            "context",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.8,
            textColor=BRAND_MUTED,
            leading=9.2,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "card_title": ParagraphStyle(
            "card_title",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.1,
            textColor=BRAND_DEEP,
            leading=10.8,
            alignment=TA_LEFT,
            spaceAfter=3,
        ),
        "card_body": ParagraphStyle(
            "card_body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.95,
            textColor=BRAND_INK,
            leading=13.0,
            alignment=TA_LEFT,
        ),
        "glance_value": ParagraphStyle(
            "glance_value",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.2,
            textColor=BRAND_DEEP,
            leading=10,
            alignment=TA_CENTER,
        ),
        "glance_note": ParagraphStyle(
            "glance_note",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.45,
            textColor=BRAND_INK,
            leading=8.85,
            alignment=TA_CENTER,
        ),
        "footer": ParagraphStyle(
            "footer",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=6.8,
            textColor=BRAND_MUTED,
            leading=8,
            alignment=TA_CENTER,
        ),
        "callout": ParagraphStyle(
            "callout",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.25,
            textColor=BRAND_INK,
            leading=9.9,
            alignment=TA_CENTER,
        ),
    }


def bullet_paragraphs(items: Iterable[str], style: ParagraphStyle) -> list[Paragraph]:
    return [Paragraph(f"- {item}", style) for item in items]


def build_pdf(variant: Variant) -> Path:
    output = OUT_DIR / f"{variant.filename_stem}.pdf"
    s = styles()
    page_width, _ = A4
    horizontal_margin = 1.0 * cm
    content_width = page_width - (2 * horizontal_margin)
    doc = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=horizontal_margin,
        leftMargin=horizontal_margin,
        topMargin=0.85 * cm,
        bottomMargin=0.7 * cm,
        title=f"SwimBuddz {variant.title} Pilot",
        author="SwimBuddz",
    )

    story = [
        Paragraph("CORPORATE WELLNESS · LAGOS", s["kicker"]),
        Paragraph(variant.audience_label, s["kicker"]),
        Paragraph(variant.title, s["title"]),
        Paragraph(variant.subtitle, s["subtitle"]),
        Paragraph(variant.intro, s["intro"]),
        Paragraph(variant.context, s["context"]),
    ]

    glance_cells = [
        [
            Paragraph(value, s["glance_value"]),
            Paragraph(note, s["glance_note"]),
        ]
        for value, note in variant.at_glance
    ]
    glance = Table(
        [glance_cells[:2], glance_cells[2:]],
        colWidths=[content_width / 2] * 2,
        rowHeights=[1.45 * cm, 1.45 * cm],
    )
    glance.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), BRAND_AQUA),
                ("BOX", (0, 0), (-1, -1), 0.6, BRAND_LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([glance, Spacer(1, 0.24 * cm)])

    section_cells = []
    for section in variant.sections:
        section_cells.append(
            [
                Paragraph(section.title.upper(), s["card_title"]),
                *bullet_paragraphs(section.bullets, s["card_body"]),
            ]
        )

    rows = [
        [section_cells[0], section_cells[1]],
        [section_cells[2], section_cells[3]],
        [section_cells[4], section_cells[5]],
    ]
    grid = Table(
        rows,
        colWidths=[content_width / 2, content_width / 2],
        rowHeights=[3.9 * cm, 3.45 * cm, 3.65 * cm],
    )
    grid.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), BRAND_BG),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF7F8")),
                ("BOX", (0, 0), (-1, -1), 0.6, BRAND_LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 11),
            ]
        )
    )
    story.extend([grid, Spacer(1, 0.24 * cm)])

    steps = [
        [Paragraph(step, s["glance_value"]), Paragraph(desc, s["glance_note"])]
        for step, desc in variant.start_steps
    ]
    step_table = Table(
        [steps[:2], steps[2:]],
        colWidths=[content_width / 2] * 2,
        rowHeights=[1.35 * cm, 1.35 * cm],
    )
    step_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.6, BRAND_LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, BRAND_LINE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend(
        [
            Paragraph("A SIMPLE FOUR-STEP START", s["card_title"]),
            step_table,
            Spacer(1, 0.18 * cm),
            Table(
                [[Paragraph(variant.next_step, s["callout"])]],
                colWidths=[content_width],
                rowHeights=[1.35 * cm],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7FCFD")),
                        ("BOX", (0, 0), (-1, -1), 0.6, BRAND_LINE),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ]
                ),
            ),
            Spacer(1, 0.12 * cm),
            Paragraph(
                "Ugo Nwachukwu · Founder, SwimBuddz · swimbuddz@gmail.com · +234 703 358 8400 · @swimbuddz · swimbuddz.com",
                s["footer"],
            ),
        ]
    )

    doc.build(story)
    return output


def add_bullets(cell, bullets: Iterable[str]) -> None:
    for item in bullets:
        p = cell.add_paragraph(style=None)
        p.style = "List Bullet"
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.add_run(item).font.size = Pt(8.2)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_cell(
    cell,
    fill: str | None = None,
    top: int = 90,
    bottom: int = 90,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=top, bottom=bottom)
    if fill:
        shade_cell(cell, fill)


def set_row_height(row, inches: float) -> None:
    row.height = Inches(inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_cell_text(cell, title: str, bullets: Iterable[str]) -> None:
    cell.text = ""
    format_cell(cell, "F5FBFC", top=200, bottom=170)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(2)
    title_p.paragraph_format.line_spacing = 1.05
    title_run = title_p.add_run(title.upper())
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(6, 59, 79)
    title_run.font.size = Pt(8.8)
    add_bullets(cell, bullets)


def build_docx(variant: Variant) -> Path:
    output = OUT_DIR / f"{variant.filename_stem}.docx"
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles_doc = document.styles
    styles_doc["Normal"].font.name = "Arial"
    styles_doc["Normal"].font.size = Pt(8.2)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(0)
    run = kicker.add_run(f"CORPORATE WELLNESS · LAGOS\n{variant.audience_label}")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(96, 114, 122)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    title_run = title.add_run(variant.title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(6, 59, 79)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle_run = subtitle.add_run(variant.subtitle)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(2)
    intro.add_run(variant.intro).font.size = Pt(8.2)

    context = document.add_paragraph()
    context.alignment = WD_ALIGN_PARAGRAPH.CENTER
    context.paragraph_format.space_after = Pt(6)
    context_run = context.add_run(variant.context)
    context_run.font.size = Pt(8)
    context_run.font.color.rgb = RGBColor(96, 114, 122)

    glance = document.add_table(rows=2, cols=2)
    glance.autofit = False
    glance.style = "Table Grid"
    for idx, (value, note) in enumerate(variant.at_glance):
        cell = glance.cell(idx // 2, idx % 2)
        format_cell(cell, "BFEFF2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = p.add_run(value + "\n")
        value_run.bold = True
        value_run.font.size = Pt(9)
        value_run.font.color.rgb = RGBColor(6, 59, 79)
        note_run = p.add_run(note)
        note_run.font.size = Pt(7)

    for row in glance.rows:
        set_row_height(row, 0.58)

    document.add_paragraph().paragraph_format.space_after = Pt(1)

    grid = document.add_table(rows=3, cols=2)
    grid.autofit = False
    grid.style = "Table Grid"
    for index, section_data in enumerate(variant.sections):
        row = index // 2
        col = index % 2
        set_cell_text(grid.cell(row, col), section_data.title, section_data.bullets)
        if row == 0:
            shade_cell(grid.cell(row, col), "EAF7F8")

    for row, height in zip(grid.rows, (1.55, 1.38, 1.46), strict=True):
        set_row_height(row, height)

    document.add_paragraph("A SIMPLE FOUR-STEP START").runs[0].bold = True
    steps = document.add_table(rows=2, cols=2)
    steps.autofit = False
    steps.style = "Table Grid"
    for idx, (step, desc) in enumerate(variant.start_steps):
        cell = steps.cell(idx // 2, idx % 2)
        format_cell(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        step_run = p.add_run(step + "\n")
        step_run.bold = True
        step_run.font.size = Pt(8.5)
        step_run.font.color.rgb = RGBColor(6, 59, 79)
        desc_run = p.add_run(desc)
        desc_run.font.size = Pt(7)

    for row in steps.rows:
        set_row_height(row, 0.54)

    callout = document.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.style = "Table Grid"
    callout_cell = callout.cell(0, 0)
    format_cell(callout_cell, "F7FCFD")
    callout_cell.text = ""
    callout_p = callout_cell.paragraphs[0]
    callout_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout_p.paragraph_format.space_after = Pt(0)
    callout_run = callout_p.add_run(variant.next_step)
    callout_run.font.size = Pt(8)
    callout_run.font.color.rgb = RGBColor(16, 42, 51)
    set_row_height(callout.rows[0], 0.54)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Ugo Nwachukwu · Founder, SwimBuddz · swimbuddz@gmail.com · +234 703 358 8400 · @swimbuddz · swimbuddz.com"
    )
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(96, 114, 122)

    document.save(output)
    return output


def main() -> None:
    for variant in (GENERIC, NESTUGE):
        pdf = build_pdf(variant)
        docx = build_docx(variant)
        print(f"Wrote {pdf}")
        print(f"Wrote {docx}")


if __name__ == "__main__":
    main()

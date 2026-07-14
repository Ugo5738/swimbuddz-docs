from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/i/Documents/work/swimbuddz/outputs/july_2026_revenue_sprint_restrategy/daniel_july_revenue_research_guide.docx"

NAVY = "0F172A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0F766E"
CYAN = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "CBD5E1"
MUTED = "475569"
ORANGE = "C2410C"
PALE_YELLOW = "FFF4CC"
PALE_GREEN = "DCFCE7"
PALE_RED = "FEE2E2"
WHITE = "FFFFFF"
BLACK = "111827"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches):
    total_dxa = sum(int(round(width * 1440)) for width in widths_inches)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = int(round(widths_inches[index] * 1440))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_inches[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def style_table(table, header_fill=CYAN, header_color=BLACK, font_size=9.2):
    table.style = "Table Grid"
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        color=header_color if row_index == 0 else BLACK,
                        bold=True if row_index == 0 else None,
                    )


def add_table(doc, headers, rows, widths, header_fill=CYAN, header_color=BLACK, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if value is None else str(value)
    set_table_geometry(table, widths)
    style_table(table, header_fill=header_fill, header_color=header_color, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def create_numbering_instance(doc, style_name="List Number", start=1):
    style = doc.styles[style_name]
    base_num_id = int(style.element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = numbering.find(f"./w:num[@w:numId='{base_num_id}']", namespaces=numbering.nsmap)
    abstract_num_id = base_num.find("w:abstractNumId", namespaces=numbering.nsmap).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall("w:num", namespaces=numbering.nsmap)]
    new_num_id = max(existing_ids) + 1

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    new_num.append(override)
    numbering.append(new_num)
    return new_num_id


def add_numbered(doc, text, num_id=None):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if num_id is not None:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, label, text, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=BLACK)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color=MID_GRAY, size="6")
    set_row_cant_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

heading_specs = {
    "Heading 1": (16, BLUE, 18, 10),
    "Heading 2": (13, BLUE, 14, 7),
    "Heading 3": (12, DARK_BLUE, 10, 5),
}
for style_name, (size, color, before, after) in heading_specs.items():
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style_name in ("List Bullet", "List Bullet 2", "List Number"):
    list_style = styles[list_style_name]
    list_style.font.name = "Calibri"
    list_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_style.font.size = Pt(11)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_paragraph.paragraph_format.space_after = Pt(0)
header_run = header_paragraph.add_run("REVENUE SPRINT GUIDE")
set_run_font(header_run, size=8.5, color=MUTED, bold=True)

footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_label = footer_paragraph.add_run("July 2026  |  ")
set_run_font(footer_label, size=9, color=MUTED)
add_page_number(footer_paragraph)

# Customer-pack style opening block.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(10)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("Immediate Cash & Swim/Sports Prospecting Guide")
set_run_font(title_run, size=25, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
subtitle_run = subtitle.add_run("Companion to the July 14-31 Revenue Research System")
set_run_font(subtitle_run, size=13, color=MUTED)

meta = add_table(
    doc,
    ["Urgent Floor", "Target", "Deadline", "Primary Cash Path"],
    [["$1,500", "$2,000", "July 31, 2026", "Warm-network tech/ops work"]],
    [1.15, 1.15, 1.55, 2.65],
    header_fill=NAVY,
    header_color=WHITE,
    font_size=9.5,
)
for cell in meta.rows[1].cells:
    set_cell_shading(cell, LIGHT_GRAY)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=10.5, bold=True, color=NAVY)

add_callout(
    doc,
    "Core decision",
    "Diversifying your search is sensible, but treating every channel equally is not. Put most of your time into trusted conversations, follow-ups, calls and paid asks. Use platforms and the swim niche as capped parallel bets.",
    fill=PALE_YELLOW,
)

add_heading(doc, "How to Use the Workbook", 1)
add_body(doc, "Open the workbook each morning and begin with the Start Here sheet. The sheets are arranged around decisions, not around information storage.")
add_table(
    doc,
    ["Sheet", "Use It For", "Required Output"],
    [
        ["Start Here", "Understand the sequence and score rules", "One next action"],
        ["Swim Research", "Qualify each swim/sports candidate", "Buyer, evidence, score and decision"],
        ["Friction Library", "Interpret what you observe", "Friction IDs and safe language"],
        ["Lead Audit", "Run discovery and sell the $400 diagnostic", "Scope, payment and next offer"],
        ["Cash Now Plan", "Protect time for the fastest cash paths", "Daily minimum outputs"],
        ["Platform Tracker", "Track applications without overinvesting", "Status and next step"],
        ["Pipeline Tracker", "Manage conversations and deals", "Every live lead has a next action"],
    ],
    [1.35, 2.55, 2.6],
    header_fill=TEAL,
    header_color=WHITE,
)

doc.add_page_break()
add_heading(doc, "1. What the Swim/Sports Work Is For", 1)
add_body(doc, "The swim/sports list is not a list of finished leads. It is a set of market anchors. A row becomes a real prospect only after you identify an operating unit, a reachable buyer, credible evidence and a reason that makes your message specific.")

add_callout(
    doc,
    "Your commercial goal",
    "Turn one public observation into a conversation, the conversation into a $400 paid audit, and a strong audit into a $1,500-$2,500 implementation sprint.",
    fill=PALE_GREEN,
)

add_heading(doc, "The Research Loop", 2)
research_numbering = create_numbering_instance(doc)
for step in [
    "Choose a specific location, franchise, school or private-lesson operator.",
    "Walk the public customer journey as a prospective parent or swimmer.",
    "Record friction evidence and separate observation from assumption.",
    "Find the owner, franchisee, general manager or operations/enrollment leader.",
    "Score the prospect in Swim Research and follow the automatic decision.",
    "Send a personalized opener only when the row is qualified.",
]:
    add_numbered(doc, step, research_numbering)

add_heading(doc, "What 'Register and Pay' Means", 2)
add_body(doc, "It means inspect how a customer would reach registration and payment. It does not mean you should buy a lesson or enter false customer details.")
add_bullet(doc, "Stop before submitting a registration form, waiver or payment.")
add_bullet(doc, "Do not provide personal information merely to continue the test.")
add_bullet(doc, "If a step is hidden behind submission, record that you could not verify it.")
add_bullet(doc, "Use language such as 'may', 'appears', or 'I could not verify' until the buyer confirms the process.")

add_heading(doc, "The 15-Minute Review", 2)
add_table(
    doc,
    ["Minutes", "Action", "Output"],
    [
        ["0-3", "Select a specific operating unit and program", "Target location/unit"],
        ["3-8", "Follow discovery, assessment, class, registration and payment pages", "Friction IDs and evidence"],
        ["8-12", "Find a buyer and contact route", "Name, role and LinkedIn/email"],
        ["12-15", "Score and choose CONTACT NOW, RESEARCH MORE or REPLACE", "One next action"],
    ],
    [0.75, 3.75, 2.0],
    header_fill=TEAL,
    header_color=WHITE,
)

add_heading(doc, "2. Qualification and Scoring", 1)
add_body(doc, "The score prevents famous brands and attractive websites from consuming your time when there is no reachable buyer or painful workflow. Use only evidence you can explain.")

add_table(
    doc,
    ["Criterion", "0", "1", "2"],
    [
        ["Reachability", "No buyer or route", "Role or generic route", "Named buyer and direct route"],
        ["Volume", "Solo, volunteer or tiny", "Small commercial operation", "Several instructors, classes or locations"],
        ["Budget", "No commercial evidence", "Commercial but uncertain", "Pricing/scale suggests ability to pay"],
        ["Pain evidence", "No friction observed", "Possible friction", "Clear observable friction"],
        ["Relevance", "Generic reason", "Segment-specific reason", "Unique observation or trigger"],
    ],
    [1.45, 1.45, 1.65, 1.95],
    header_fill=TEAL,
    header_color=WHITE,
)

add_heading(doc, "Decision Rules", 2)
decision_table = add_table(
    doc,
    ["Score", "Decision", "Action"],
    [
        ["8-10", "CONTACT NOW", "Send the personalized opener today. Reachability and pain must each be at least 1."],
        ["6-7", "RESEARCH MORE", "Spend no more than 10 additional minutes finding the buyer or stronger evidence."],
        ["0-5", "REPLACE", "Replace the company with a smaller or more reachable operator."],
    ],
    [0.8, 1.35, 4.35],
    header_fill=ORANGE,
    header_color=WHITE,
)
set_cell_shading(decision_table.rows[1].cells[1], PALE_GREEN)
set_cell_shading(decision_table.rows[2].cells[1], PALE_YELLOW)
set_cell_shading(decision_table.rows[3].cells[1], PALE_RED)

add_heading(doc, "Evidence Examples", 2)
add_bullet(doc, "Volume: several locations, many weekly classes, multiple instructors, waitlists or active hiring.")
add_bullet(doc, "Budget: published commercial pricing, franchise operation, paid staff, several programs or expansion activity.")
add_bullet(doc, "Reachability: a named owner/operator with LinkedIn, email, contact form or warm introduction.")
add_bullet(doc, "Pain: an observable form, portal or handoff that creates delay, duplicate work or uncertainty.")
add_bullet(doc, "Relevance: one observation that could only have been written after reviewing this exact business.")

doc.add_page_break()
add_heading(doc, "Hard Disqualifiers", 2)
add_bullet(doc, "Volunteer-only club with no commercial buyer.")
add_bullet(doc, "No reachable owner, operator or route to one.")
add_bullet(doc, "No repeated enrollment or operating workflow.")
add_bullet(doc, "No plausible ability to pay $400.")
add_bullet(doc, "A mature system already handles the identified workflow well.")

add_heading(doc, "3. Website and Workflow Friction", 1)
add_body(doc, "Friction is evidence that a customer or staff member must think, wait, re-enter information or depend on a manual handoff. It is not permission to claim lost revenue. The Friction Library contains the full list and safe wording.")

add_table(
    doc,
    ["Stage", "Common Friction", "What to Record", "Safe Interpretation"],
    [
        ["Discovery", "Unclear CTA, hidden pricing or availability", "Page, button text, missing information", "The first step may require extra explanation."],
        ["Inquiry", "Generic form with little routing context", "Fields and confirmation behavior", "The team may need to re-qualify each inquiry."],
        ["Assessment", "Request submitted without slot selection", "Calendar or phone/email handoff", "There may be a manual scheduling step."],
        ["Placement", "Unclear level, location or instructor matching", "Level/location guidance and routing", "Staff judgment may be repeated for each lead."],
        ["Registration", "Repeated data, separate waiver or portal", "Portal changes and duplicate fields", "The journey appears to repeat information."],
        ["Payment", "Invoice later or separate checkout", "Payment timing and system", "Payment may be a separate operational handoff."],
        ["Follow-up", "Unclear confirmation, response time or next owner", "Success message and stated timing", "The customer may not know what happens next."],
        ["Operations", "Waitlist, rescheduling or matching handled manually", "Self-service options and discovery answers", "Routine changes may require staff intervention."],
    ],
    [1.0, 2.05, 1.75, 1.7],
    header_fill=TEAL,
    header_color=WHITE,
    font_size=8.8,
)

add_heading(doc, "Scoring Friction", 2)
add_bullet(doc, "0: no issue observed or the step appears clear and self-service.")
add_bullet(doc, "1: possible friction, uncertainty, or a step you cannot verify publicly.")
add_bullet(doc, "2: a clear observable handoff, delay, duplicate step or missing path.")
add_body(doc, "A high friction total is not enough by itself. A prospect still needs a reachable buyer, commercial volume and a personalized reason.")

add_callout(
    doc,
    "Strong pain hypothesis",
    "The assessment request ends in a generic form before a customer can select a time. That may create a manual follow-up step between inquiry and assessment booking.",
    fill=PALE_GREEN,
)
add_callout(
    doc,
    "Weak claim",
    "You are losing lots of customers because your website is bad. This is unsupported, adversarial and unlikely to create trust.",
    fill=PALE_RED,
)

add_heading(doc, "4. From Research to a Paid Conversation", 1)
add_body(doc, "The first message should earn permission to continue. It should not attempt to sell a complex automation system to a stranger.")

add_heading(doc, "Cold Opener", 2)
add_callout(
    doc,
    "Template",
    "Hi [Name], I run SwimBuddz and build backend/workflow systems. I reviewed the public enrollment path for [location] and noticed [specific observation]. There may be a manual handoff between [step] and [step]. Would it be useful if I sent one short observation showing what I mean?",
    fill=LIGHT_GRAY,
)

add_heading(doc, "After Interest", 2)
add_callout(
    doc,
    "Template",
    "There may be two or three related points in the full workflow. Before recommending automation, I would want to understand how your team handles inquiries, assessments and follow-up. Are you open to a 15-minute call this week?",
    fill=LIGHT_GRAY,
)

add_heading(doc, "What You Give Away", 2)
add_bullet(doc, "One specific observation or a two-to-three minute Loom.")
add_bullet(doc, "Enough context to show that you understand the workflow.")
add_bullet(doc, "A clear invitation to a 15-minute discovery call.")

add_heading(doc, "What Remains Paid", 2)
add_bullet(doc, "The full journey map.")
add_bullet(doc, "The complete friction and lead-leak assessment.")
add_bullet(doc, "Prioritization, quick wins and implementation blueprint.")
add_bullet(doc, "The review call and scoped build recommendation.")

add_heading(doc, "Follow-up", 2)
add_body(doc, "Follow up after two business days with a question that can produce a useful answer:")
add_callout(
    doc,
    "Template",
    "Quick follow-up: is [specific workflow] currently a real operational problem for your team, or is it already handled well enough?",
    fill=PALE_YELLOW,
)
add_body(doc, "After one or two concise follow-ups, close the loop and replace the prospect. Silence is not a reason to keep researching the company.")

add_heading(doc, "5. The $400 Lead-to-Enrollment Leak Audit", 1)
add_body(doc, "The audit is a paid diagnostic, not a disguised sales call. It should help the buyer make a build/no-build decision even if they do not hire you for implementation.")

add_table(
    doc,
    ["Commercial Term", "Definition"],
    [
        ["Price", "$400, paid upfront"],
        ["Delivery", "Within 48 hours after discovery and required inputs"],
        ["Credit", "$400 credited toward a sprint started within 30 days"],
        ["Scope", "One lead-to-enrollment workflow"],
        ["Next offer", "$1,500-$2,500 implementation sprint with 50% deposit"],
    ],
    [1.55, 4.95],
    header_fill=TEAL,
    header_color=WHITE,
)

add_heading(doc, "Audit Deliverables", 2)
for deliverable in [
    "Current journey map showing systems, owners and handoffs.",
    "Evidence-based friction list with observations separated from assumptions.",
    "Prioritized assessment using impact, confidence and effort.",
    "Three quick wins that do not require a full build.",
    "Implementation blueprint for one recommended workflow.",
    "Thirty-minute review call and sprint decision.",
]:
    add_bullet(doc, deliverable)

add_heading(doc, "Paid Audit Close", 2)
add_callout(
    doc,
    "Template",
    "Based on what you described, I would start with a 48-hour Lead-to-Enrollment Leak Audit. I will map the current journey, identify the main leak and administrative points, recommend quick wins, and give you a scoped implementation plan. It is $400 paid upfront, and I credit it toward a sprint started within 30 days. Shall I send the scope and payment link?",
    fill=PALE_GREEN,
)

add_heading(doc, "Do Not Start Until", 2)
add_bullet(doc, "The buyer has agreed to the written scope.")
add_bullet(doc, "Payment is received.")
add_bullet(doc, "The discovery call is complete.")
add_bullet(doc, "Required access or screenshots are available.")

add_heading(doc, "6. Immediate Cash Strategy", 1)
add_body(doc, "Your overall thinking is directionally right: combine immediate work, direct client acquisition and longer-term channels. The mistake would be giving equal time to channels with very different time-to-cash.")

add_table(
    doc,
    ["Path", "Time Share", "Daily Minimum", "Role"],
    [
        ["Warm-network tech/ops", "50%", "10 direct asks/follow-ups plus active calls", "Primary cash path"],
        ["Agency + high-fit platforms", "25%", "2 agency asks and 3 strong applications", "Parallel cash"],
        ["Swim/sports niche", "15%", "2 qualified reviews and 2 messages", "Controlled market test"],
        ["Corporate wellness/cohorts", "10%", "Warm follow-ups and existing leads only", "Backup"],
        ["StrokeLab", "0% build time", "Use as proof", "Authority asset"],
    ],
    [2.25, 0.9, 2.25, 1.1],
    header_fill=TEAL,
    header_color=WHITE,
)

add_heading(doc, "Platforms From Nigeria", 2)
add_body(doc, "Use platforms as a capped parallel channel, not as money you can already count. Application, screening, task availability, client approval and payout timing can all delay cash.")

platform_table = add_table(
    doc,
    ["Platform", "Use", "Caution"],
    [
        ["Outlier", "Apply once to a genuinely global coding/AI role", "Eligibility and tasks are location/project dependent. Use your real Nigeria location."],
        ["Upwork", "Submit three highly matched fixed-scope proposals daily", "New profiles face competition; fixed-price payout includes a security period."],
        ["Andela", "Complete one quality application and assessment", "Good medium-term route, not reliable for immediate July cash."],
        ["Contra", "Publish the audit and sprint as packaged services", "Better as a storefront/direct-contract layer than an instant job feed."],
    ],
    [1.0, 2.6, 2.9],
    header_fill=TEAL,
    header_color=WHITE,
)

add_callout(
    doc,
    "VPN rule",
    "Do not use a VPN to pretend you live in an eligible country. Outlier says contributors may work only from their verified primary location and accounts may be paused when location cannot be verified. Use Nigeria truthfully and apply only where the role accepts it.",
    fill=PALE_RED,
)

add_heading(doc, "7. Daily Operating Playbook", 1)
add_body(doc, "Do these blocks in order. When time is short, protect the first three because they are closest to cash.")
add_table(
    doc,
    ["Order", "Block", "Time Cap", "Output"],
    [
        [1, "Replies and follow-ups", "45 min", "Every live lead has a call, paid ask, proposal or clear no"],
        [2, "Warm direct asks", "75 min", "10 quality touches or referral asks"],
        [3, "Calls, proposals and delivery", "90 min", "Cash-stage movement"],
        [4, "Platforms and agencies", "60-90 min", "3 matched applications and 2 agency asks"],
        [5, "Swim niche", "30-45 min", "2 completed reviews and 2 personalized messages"],
        [6, "Daily score", "15 min", "Replies, calls, proposals, cash and market signal recorded"],
    ],
    [0.65, 2.1, 1.05, 2.7],
    header_fill=ORANGE,
    header_color=WHITE,
)

add_heading(doc, "Your First 90 Minutes", 2)
first_90_numbering = create_numbering_instance(doc)
for step in [
    "Open Pipeline Tracker and handle every existing reply or overdue follow-up.",
    "Send five direct messages to people who know your work and can buy or introduce you.",
    "Open Swim Research and complete one company row from location through decision.",
    "If the row says CONTACT NOW, send the opener and schedule its follow-up.",
    "Apply to one highly matched platform role only after those cash conversations are moving.",
]:
    add_numbered(doc, step, first_90_numbering)

add_heading(doc, "Daily Questions", 2)
add_bullet(doc, "How many real conversations did I create today?")
add_bullet(doc, "Which conversations moved toward a call, proposal or payment?")
add_bullet(doc, "What did the market reject, question or respond to?")
add_bullet(doc, "Which channel deserves more time tomorrow?")
add_bullet(doc, "What can I stop doing because it is not moving toward cash?")

add_callout(
    doc,
    "Final rule",
    "Do not measure the day by websites reviewed, profiles created or messages drafted. Measure replies, calls, paid asks, proposals, deposits and useful market feedback.",
    fill=PALE_GREEN,
)

add_heading(doc, "Official Platform Sources", 3)
sources = [
    ("Outlier Working Location Policy", "https://outlier.ai/legal/working-location-policy"),
    ("Upwork Nigeria VAT guidance", "https://support.upwork.com/hc/en-us/articles/15156139408147-How-VAT-works-for-freelancers-in-Nigeria"),
    ("Upwork payout timing", "https://support.upwork.com/hc/en-us/articles/211060918-Manage-How-You-Get-Paid"),
    ("Andela approved regions", "https://help.andela.com/hc/en-us/articles/32941472534035-Where-must-I-live-to-work-at-Andela"),
    ("Andela application process", "https://help.andela.com/hc/en-us/articles/26424660453395-How-can-I-join-the-Andela-Talent-Cloud"),
    ("Contra global payments", "https://contra.com/features/global-payments"),
]
for label, url in sources:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(2)
    add_hyperlink(paragraph, label, url)

doc.core_properties.title = "Immediate Cash & Swim/Sports Prospecting Guide"
doc.core_properties.subject = "July 2026 revenue execution and swim/sports prospect qualification"
doc.core_properties.author = "Daniel"
doc.core_properties.keywords = "revenue, prospecting, swim school, paid audit, tech gigs"
doc.save(OUTPUT)
print(OUTPUT)
